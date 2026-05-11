"""Генератор boilerplate-агента из wizard-спеки.

Юзер в UI заполняет id/name/inputs/outputs — мы генерим минимальный скелет
(manifest.yaml + agent.py + requirements.txt + Dockerfile) и регистрируем
через stage_template_as_local_repo.
"""
from __future__ import annotations

from dataclasses import dataclass
from typing import Literal

import yaml

InputType = Literal["text", "textarea", "number", "checkbox", "select", "radio"]
OutputType = Literal["docx", "pdf", "json", "txt", "any"]


@dataclass
class TemplateInput:
    id: str
    type: InputType
    label: str
    required: bool = False


@dataclass
class TemplateOutput:
    id: str
    type: OutputType
    label: str
    filename: str
    primary: bool = False


@dataclass
class TemplateSpec:
    slug: str  # id агента (slug)
    name: str
    icon: str | None
    category: str  # 'научная-работа' / 'учебная' / 'организационная'
    short_description: str
    inputs: list[TemplateInput]
    outputs: list[TemplateOutput]
    use_llm: bool = False


def _manifest_yaml(spec: TemplateSpec) -> str:
    manifest = {
        "id": spec.slug,
        "name": spec.name,
        "version": "0.1.0",
        "category": spec.category,
        "short_description": spec.short_description.strip() + "\n",
    }
    if spec.icon:
        manifest["icon"] = spec.icon

    inputs_dict = {}
    for inp in spec.inputs:
        block: dict = {"type": inp.type, "label": inp.label}
        if inp.required:
            block["required"] = True
        inputs_dict[inp.id] = block
    if inputs_dict:
        manifest["inputs"] = inputs_dict

    manifest["files"] = {}

    outputs = []
    for out in spec.outputs:
        block = {
            "id": out.id, "type": out.type,
            "label": out.label, "filename": out.filename,
        }
        if out.primary:
            block["primary"] = True
        outputs.append(block)
    manifest["outputs"] = outputs

    runtime = {
        "docker": {
            "base_image": "python:3.12-slim",
            "setup": ["pip install -r requirements.txt"],
            "entrypoint": ["python", "agent.py"],
        },
        "limits": {
            "max_runtime_minutes": 10,
            "max_memory_mb": 512,
            "max_cpu_cores": 1,
        },
    }
    if spec.use_llm:
        runtime["llm"] = {
            "provider": "openrouter",
            "models": ["anthropic/claude-3.5-haiku", "deepseek/deepseek-r1"],
        }
    manifest["runtime"] = runtime
    return yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False)


def _agent_py(spec: TemplateSpec) -> str:
    primary = next((o for o in spec.outputs if o.primary), spec.outputs[0] if spec.outputs else None)
    input_reads = "\n".join(
        f"    {inp.id} = params.get({inp.id!r})" for inp in spec.inputs
    ) or "    # параметров нет"

    llm_block = ""
    if spec.use_llm:
        llm_block = '''
    # Пример LLM-вызова через прокси портала.
    # api_key = os.environ["OPENROUTER_API_KEY"]
    # base_url = os.environ["OPENROUTER_BASE_URL"]
    # import httpx
    # r = httpx.post(
    #     f"{base_url}/chat/completions",
    #     headers={"Authorization": f"Bearer {api_key}"},
    #     json={"model": "anthropic/claude-3.5-haiku",
    #           "messages": [{"role": "user", "content": "Привет!"}]},
    #     timeout=120,
    # )
    # text = r.json()["choices"][0]["message"]["content"]
'''

    primary_path = primary.filename if primary else "result.txt"
    primary_id = primary.id if primary else "result"
    primary_type = primary.type if primary else "txt"

    body = _output_writer(primary_type, primary_path)
    artifacts = ",\n        ".join(
        f'{{"id": {o.id!r}, "path": {o.filename!r}}}' for o in spec.outputs
    )

    return f'''"""{spec.name} — сгенерировано wizard'ом mirea-agent-portal.

Замени логику ниже на свою. Документация SDK:
https://github.com/svpeditor/mirea-agent-portal/blob/main/docs/agent-developer-guide.md
"""
from __future__ import annotations

import json
import os
from pathlib import Path

from portal_sdk import Agent


def main() -> None:
    agent = Agent()
    params = agent.params
{input_reads}

    agent.log("info", f"Запуск {spec.name!r} с params={{params}}")
    agent.progress(0.1, "Начинаем работу")
{llm_block}
    # TODO: тут твоя бизнес-логика.
    agent.progress(0.5, "Обработка")
{body}
    agent.progress(1.0, "Готово")
    agent.result(artifacts=[
        {artifacts}
    ])


if __name__ == "__main__":
    main()
'''


def _output_writer(otype: str, filename: str) -> str:
    if otype == "docx":
        return f'''    out_dir = agent.output_dir
    from docx import Document
    doc = Document()
    doc.add_heading("{filename}", level=0)
    doc.add_paragraph("Результат работы агента.")
    doc.save(out_dir / "{filename}")
'''
    if otype == "json":
        return f'''    out_dir = agent.output_dir
    (out_dir / "{filename}").write_text(
        json.dumps({{"ok": True, "params": params}}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
'''
    if otype in ("txt", "any"):
        return f'''    out_dir = agent.output_dir
    (out_dir / "{filename}").write_text(
        f"Параметры: {{params}}\\n",
        encoding="utf-8",
    )
'''
    return f'''    out_dir = agent.output_dir
    (out_dir / "{filename}").write_bytes(b"binary result")
'''


def _requirements(spec: TemplateSpec) -> str:
    lines = []
    has_docx = any(o.type == "docx" for o in spec.outputs)
    if has_docx:
        lines.append("python-docx>=1.1,<2")
    if spec.use_llm:
        lines.append("httpx>=0.27,<1")
    return "\n".join(lines) + ("\n" if lines else "")


def _dockerfile() -> str:
    return """FROM python:3.12-slim

WORKDIR /agent

COPY requirements.txt ./
RUN pip install --no-cache-dir -r requirements.txt

COPY agent.py manifest.yaml ./

ENTRYPOINT ["python", "agent.py"]
"""


def _readme(spec: TemplateSpec) -> str:
    return f"""# {spec.name}

{spec.short_description}

Сгенерировано wizard'ом mirea-agent-portal.

## Изменить логику

1. Открой `agent.py`, замени TODO-секцию на свою.
2. При необходимости добавь зависимости в `requirements.txt`.
3. Обнови `manifest.yaml`: inputs / outputs / runtime.
4. Загрузи новую версию через admin-портал.

См. полную документацию: https://github.com/svpeditor/mirea-agent-portal/blob/main/docs/agent-developer-guide.md
"""


def build_template_files(spec: TemplateSpec) -> dict[str, str]:
    return {
        "manifest.yaml": _manifest_yaml(spec),
        "agent.py": _agent_py(spec),
        "requirements.txt": _requirements(spec),
        "Dockerfile": _dockerfile(),
        "README.md": _readme(spec),
    }
