"""science_agent_stub — заглушка под контракт SDK v0.1.

Stub-уровень: имитирует поиск научных статей. Возвращает фиксированный
mock-набор «статей», подсчитывает прогресс, формирует Word-отчёт и BibTeX.

Дополнительно: smoke-test 1.2.4-обвязки — логирует наличие
OPENROUTER_API_KEY (ephemeral-токен) в env, не вызывая LLM. Real-версия
подтянется в milestone-1 (см. docs/superpowers/plans/2026-05-10-1.4-agents-port.md).
"""
from __future__ import annotations

import os
from dataclasses import dataclass

from docx import Document

from portal_sdk import Agent


@dataclass(frozen=True)
class MockPaper:
    paper_id: str
    title: str
    authors: str
    venue: str
    year: int


MOCK_PAPERS: list[MockPaper] = [
    MockPaper("arxiv:2401.01234", "Self-Supervised Learning at Scale",
              "Liu et al.", "ICLR", 2024),
    MockPaper("arxiv:2403.05678", "Sparse Mixture-of-Experts Revisited",
              "Chen, Park", "NeurIPS", 2024),
    MockPaper("doi:10.1145/123456", "Energy-Efficient Transformers",
              "Müller, Sato, Khan", "ACM TOPLAS", 2023),
    MockPaper("arxiv:2410.00910", "Long-Context Reasoning Without RAG",
              "Iyengar, Wang", "EMNLP", 2024),
    MockPaper("arxiv:2502.04321", "Differentially Private Fine-Tuning of LLMs",
              "Abadi, McMahan", "USENIX Security", 2025),
    MockPaper("arxiv:2406.17234", "Robust Vision-Language Alignment",
              "Patel, Gomez", "CVPR", 2024),
    MockPaper("arxiv:2308.09988", "Scaling Laws Beyond Cross-Entropy",
              "Hoffmann et al.", "ICML", 2023),
    MockPaper("doi:10.1109/TPAMI.2024.55", "Attention Is All You Need (10-Year Retrospective)",
              "Vaswani, Polosukhin", "TPAMI", 2024),
    MockPaper("arxiv:2503.14159", "Cooperative Multi-Agent Reasoning",
              "Singh, Tan", "AAAI", 2025),
    MockPaper("arxiv:2412.07654", "Federated Learning at Hospital Scale",
              "Kim, García", "JAMIA", 2024),
]


def _bibtex_entry(p: MockPaper) -> str:
    key = p.paper_id.replace(":", "_").replace(".", "_").replace("/", "_")
    return (
        f"@article{{{key},\n"
        f"  title = {{{p.title}}},\n"
        f"  author = {{{p.authors}}},\n"
        f"  journal = {{{p.venue}}},\n"
        f"  year = {{{p.year}}}\n"
        "}\n"
    )


def main() -> None:
    agent = Agent()
    params = agent.params
    topic: str = params.get("topic", "(не указано)")
    max_papers: int = int(params.get("max_papers", 30))
    language: str = params.get("language", "en")

    has_token = bool(os.environ.get("OPENROUTER_API_KEY"))
    agent.log(
        "info",
        f"science_agent_stub: topic={topic[:60]!r}, max_papers={max_papers}, "
        f"language={language}, llm_token_present={has_token}",
    )

    # Stub: возвращаем min(max_papers, len(MOCK_PAPERS)) статей.
    selected = MOCK_PAPERS[: min(max_papers, len(MOCK_PAPERS))]
    n = len(selected)
    if n == 0:
        agent.failed("Mock-набор пуст — это баг stub-агента.")
        return

    for i, paper in enumerate(selected):
        agent.progress((i + 1) / n, f"Анализ {i + 1}/{n}: {paper.title}")
        agent.item_done(
            paper.paper_id,
            summary=paper.title,
            data={"venue": paper.venue, "year": paper.year, "score": round(0.95 - 0.05 * i, 2)},
        )

    out_dir = agent.output_dir

    report = Document()
    report.add_heading("Поиск научных статей — отчёт", level=0)
    report.add_paragraph(f"Тема: {topic}")
    report.add_paragraph(f"Язык запроса: {language}")
    report.add_paragraph(f"Найдено статей (mock): {n}")
    report.add_heading("Ранжированный список", level=1)
    for i, paper in enumerate(selected, start=1):
        report.add_heading(f"{i}. {paper.title}", level=2)
        report.add_paragraph(f"Авторы: {paper.authors}")
        report.add_paragraph(f"Источник: {paper.venue} ({paper.year})")
        report.add_paragraph(f"Идентификатор: {paper.paper_id}")
        report.add_paragraph(
            f"Mock-аннотация: статья релевантна теме «{topic[:60]}»; разбор будет в real-версии.",
        )
    report.add_paragraph()
    report.add_paragraph("Это отчёт от stub-агента. Real-версия использует LLM-ранжирование.")
    report.save(out_dir / "report.docx")

    bib = "".join(_bibtex_entry(p) for p in selected)
    (out_dir / "sources.bib").write_text(bib, encoding="utf-8")

    agent.result(artifacts=[
        {"id": "report", "path": "report.docx"},
        {"id": "bibtex", "path": "sources.bib"},
    ])


if __name__ == "__main__":
    main()
