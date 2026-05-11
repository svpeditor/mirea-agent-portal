"""proverka_stub — заглушка под контракт SDK v0.1.

Stub-уровень: имитирует анализ работ школьников. Сканирует структуру
$INPUT_DIR/works/<подпапка>, считает файлы и размер, возвращает
сводный Word и zip per-work-заглушек. Mock-оценка случайная.

Real-версия живёт у Дани и подтянется в milestone-1 этого плана
(см. docs/superpowers/plans/2026-05-10-1.4-agents-port.md).
"""
from __future__ import annotations

import io
import random
import zipfile
from pathlib import Path

from docx import Document

from portal_sdk import Agent

CHECKLIST = [
    "Соответствие теме",
    "Научная новизна",
    "Качество эксперимента",
    "Оформление",
]


def _collect_files(folder: Path) -> tuple[int, int]:
    total = 0
    size = 0
    for p in folder.rglob("*"):
        if p.is_file():
            total += 1
            size += p.stat().st_size
    return total, size


def _per_work_doc(work_name: str, files_n: int, size_kb: int, score: int) -> bytes:
    doc = Document()
    doc.add_heading(f"Заключение: {work_name}", level=1)
    doc.add_paragraph(f"Файлов в работе: {files_n}")
    doc.add_paragraph(f"Объём: {size_kb} КБ")
    doc.add_paragraph(f"Mock-оценка: {score} из 100")
    doc.add_paragraph("Это stub-заключение. Real-версия будет давать развёрнутый разбор по чек-листу.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> None:
    agent = Agent()
    params = agent.params
    competition: str = params.get("competition", "(без названия)")
    grade_level: str = params.get("grade_level", "10-11")

    agent.log("info", f"proverka_stub: конкурс={competition!r}, класс={grade_level!r}")

    works_dir = agent.input_dir("works")
    work_folders = sorted(p for p in works_dir.iterdir() if p.is_dir())
    n = len(work_folders)
    if n == 0:
        agent.failed("В папке works нет ни одной подпапки-работы.")
        return

    rng = random.Random(42)  # детерминированно для stub-проверок
    rows: list[tuple[str, int, int, int]] = []

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, folder in enumerate(work_folders):
            agent.progress((i + 1) / n, f"Работа {i + 1} из {n}: {folder.name}")
            files_n, size_b = _collect_files(folder)
            size_kb = size_b // 1024
            score = rng.randint(50, 95)
            rows.append((folder.name, files_n, size_kb, score))
            agent.item_done(
                folder.name,
                summary=f"{files_n} файлов, {size_kb} КБ, оценка {score}",
                data={"files": files_n, "size_kb": size_kb, "score": score},
            )
            zf.writestr(
                f"{folder.name}.docx",
                _per_work_doc(folder.name, files_n, size_kb, score),
            )

    out_dir = agent.output_dir
    (out_dir / "per_work.zip").write_bytes(zip_buf.getvalue())

    report = Document()
    report.add_heading(f"{competition} — сводное заключение", level=0)
    report.add_paragraph(f"Класс участников: {grade_level}")
    report.add_paragraph(f"Работ принято к проверке: {n}")
    report.add_heading("Чек-лист научной экспертизы", level=1)
    for item in CHECKLIST:
        report.add_paragraph(f"• {item}", style="List Bullet")
    report.add_heading("Результаты", level=1)
    table = report.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Работа"
    hdr[1].text = "Файлов"
    hdr[2].text = "Размер, КБ"
    hdr[3].text = "Оценка"
    for name, files_n, size_kb, score in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = str(files_n)
        row[2].text = str(size_kb)
        row[3].text = str(score)
    report.add_paragraph()
    report.add_paragraph(
        "Это заключение сформировано stub-агентом. Real-версия подтянется в milestone-1.",
    )
    report.save(out_dir / "report.docx")

    agent.result(artifacts=[
        {"id": "report", "path": "report.docx"},
        {"id": "per_work", "path": "per_work.zip"},
    ])


if __name__ == "__main__":
    main()
