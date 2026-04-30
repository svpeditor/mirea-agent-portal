"""Echo-агент. Самый простой пример работы с portal_sdk."""
from __future__ import annotations

import json
import time

from docx import Document

from portal_sdk import Agent


def main() -> None:
    agent = Agent()
    params = agent.params

    message: str = params["message"]
    loops: int = int(params.get("loops", 5))
    shout: bool = bool(params.get("shout", False))

    if shout:
        message = message.upper()

    agent.log("info", f"Эхо: '{message}' × {loops}")

    doc = Document()
    doc.add_heading("Echo", level=0)
    doc.add_paragraph(f"Параметры: loops={loops}, shout={shout}")

    for i in range(loops):
        agent.progress((i + 1) / loops, f"Строка {i + 1} из {loops}")
        doc.add_paragraph(f"{i + 1}. {message}")
        agent.item_done(f"line-{i + 1}", summary=f"добавлена строка {i + 1}")
        time.sleep(0.1)  # чуть задержки чтобы прогресс был видим

    out_dir = agent.output_dir
    doc.save(out_dir / "echo.docx")

    (out_dir / "summary.json").write_text(
        json.dumps(
            {"message": message, "loops": loops, "shout": shout},
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    agent.result(artifacts=[
        {"id": "report", "path": "echo.docx"},
        {"id": "summary", "path": "summary.json"},
    ])


if __name__ == "__main__":
    main()
